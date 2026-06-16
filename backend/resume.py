"""Resume Optimizer.

Takes a job description + an existing Word resume, asks the Claude model (via the Agent
SDK, OAuth only) to rewrite/optimize the resume points to match the requirement, and
produces a downloadable Word document.

State is persisted per user in the `resume_jobs` table so a browser refresh mid-generation
restores the progress bar and result.
"""

import io
import json
import os
import re
import tempfile
from datetime import datetime, timezone

# Repo root holds `.claude/skills/` — the SDK discovers project skills relative to cwd.
REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt
from fastapi import (
    APIRouter,
    BackgroundTasks,
    Body,
    Depends,
    File,
    Form,
    HTTPException,
    UploadFile,
)
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# Same OAuth-only rule as the rest of the backend: never authenticate via an API key.
os.environ.pop("ANTHROPIC_API_KEY", None)
os.environ.pop("ANTHROPIC_AUTH_TOKEN", None)

from claude_agent_sdk import ClaudeSDKClient, ClaudeAgentOptions  # noqa: E402
from claude_agent_sdk.types import (  # noqa: E402
    AssistantMessage,
    TextBlock,
    ResultMessage,
)

from db import get_db_connection  # noqa: E402
from auth import get_current_user  # noqa: E402


# ---------------------------------------------------------------------------
# Schema
# ---------------------------------------------------------------------------
def init_resume_db():
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS resume_jobs (
            user_id INTEGER PRIMARY KEY,
            status TEXT DEFAULT 'idle',
            stage TEXT DEFAULT '',
            progress INTEGER DEFAULT 0,
            job_description TEXT DEFAULT '',
            original_text TEXT DEFAULT '',
            result_markdown TEXT DEFAULT '',
            result_json TEXT DEFAULT '',
            result_docx BLOB,
            error TEXT DEFAULT '',
            updated_at TEXT
        )
        """
    )
    # Migrate older databases that predate the structured result_json column.
    cur.execute("PRAGMA table_info(resume_jobs)")
    cols = {row[1] for row in cur.fetchall()}
    if "result_json" not in cols:
        cur.execute("ALTER TABLE resume_jobs ADD COLUMN result_json TEXT DEFAULT ''")
    conn.commit()
    conn.close()


def _set_job(user_id: int, **fields):
    fields["updated_at"] = datetime.now(timezone.utc).isoformat()
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT user_id FROM resume_jobs WHERE user_id = ?", (user_id,))
    exists = cur.fetchone() is not None
    if exists:
        set_clause = ", ".join(f"{k} = ?" for k in fields)
        cur.execute(
            f"UPDATE resume_jobs SET {set_clause} WHERE user_id = ?",
            (*fields.values(), user_id),
        )
    else:
        cols = ", ".join(["user_id", *fields.keys()])
        placeholders = ", ".join(["?"] * (len(fields) + 1))
        cur.execute(
            f"INSERT INTO resume_jobs ({cols}) VALUES ({placeholders})",
            (user_id, *fields.values()),
        )
    conn.commit()
    conn.close()


def _get_job(user_id: int):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT * FROM resume_jobs WHERE user_id = ?", (user_id,))
    row = cur.fetchone()
    conn.close()
    return row


# ---------------------------------------------------------------------------
# Structured content model
#
# The optimizer returns a structured resume so the UI can show a diff (which
# bullets are newly added) and let the user edit each point. `is_new` marks
# bullets the optimizer added to match the JD; original bullets are preserved.
# ---------------------------------------------------------------------------
class ResumeItem(BaseModel):
    text: str
    is_new: bool = False


class ResumeSection(BaseModel):
    title: str
    items: list[ResumeItem] = []


class ResumeContent(BaseModel):
    summary: str = ""
    sections: list[ResumeSection] = []


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------
def extract_text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]
    # Include table cell text (resumes sometimes use tables for layout).
    for table in doc.tables:
        for r in table.rows:
            for c in r.cells:
                if c.text.strip():
                    lines.append(c.text)
    return "\n".join(lines).strip()


# Leading bullet glyphs/markers a line might already carry (from the original resume
# text). They must be stripped before applying a real "List Bullet" style, otherwise
# the document shows a double bullet (the style's bullet + the literal glyph).
_BULLET_PREFIX = re.compile(r"^\s*[•‣◦⁃∙·\*\-–—o]\s+")


def strip_bullet(text: str) -> str:
    """Remove a single leading bullet glyph/marker from a line of text."""
    return _BULLET_PREFIX.sub("", (text or "").strip(), count=1).strip()


def content_to_markdown(content: dict) -> str:
    """Derive a simple markdown form (used for preview/fallback)."""
    lines = []
    summary = (content.get("summary") or "").strip()
    if summary:
        lines.append("## Professional Summary")
        lines.append(summary)
    for section in content.get("sections", []):
        title = (section.get("title") or "").strip()
        if title:
            lines.append(f"## {title}")
        for item in section.get("items", []):
            text = strip_bullet(item.get("text") or "")
            if text:
                lines.append(f"- {text}")
    return "\n".join(lines)


def build_docx_from_content(content: dict) -> bytes:
    """Build a clean, professional .docx from structured resume content.

    Follows the `docx` skill's guidance: explicit US-Letter page size, Arial body
    font, proper heading styles, and real bulleted lists (no unicode bullets). The
    downloaded document is intentionally clean — the new-vs-original diff is a UI
    concern and is not baked into the file.
    """
    doc = Document()

    # US Letter page size with 1" margins (skill: docx defaults to A4 — set explicitly).
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    summary = (content.get("summary") or "").strip()
    if summary:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(summary)

    for sec in content.get("sections", []):
        title = (sec.get("title") or "").strip()
        if title:
            doc.add_heading(title, level=1)
        for item in sec.get("items", []):
            text = strip_bullet(item.get("text") or "")
            if text:
                doc.add_paragraph(text, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Claude call
# ---------------------------------------------------------------------------
def _parse_content(text: str) -> dict:
    """Parse Claude's response into the structured content dict, tolerating fences."""
    match = re.search(r"```json\s*(\{.*\})\s*```", text, re.DOTALL | re.IGNORECASE)
    if not match:
        match = re.search(r"(\{.*\})", text, re.DOTALL)
    raw = match.group(1) if match else text
    data = json.loads(raw)
    # Coerce/validate via the pydantic model so the shape is guaranteed.
    return sanitize_content(ResumeContent(**data).model_dump())


def sanitize_content(content: dict) -> dict:
    """Strip any leading bullet glyph each item already carries, so neither the
    structured editor nor the generated .docx ends up with a double bullet."""
    for section in content.get("sections", []):
        for item in section.get("items", []):
            item["text"] = strip_bullet(item.get("text") or "")
    return content


async def _optimize_with_claude(job_description: str, resume_text: str) -> tuple[dict, bytes | None]:
    """Optimize the resume with Claude via the Agent SDK.

    The bundled **docx skill** (`.claude/skills/docx/`) is enabled on the SDK options so the
    agent can produce a polished Word document directly. The agent returns structured content
    (for the UI diff) AND, when possible, writes the optimized `.docx` to `out_path`. If the
    skill path is unavailable, callers fall back to `build_docx_from_content`.

    Returns `(content_dict, docx_bytes_or_None)`.
    """
    out_fd, out_path = tempfile.mkstemp(suffix=".docx", prefix="optimized_resume_")
    os.close(out_fd)
    os.remove(out_path)  # let the agent create it fresh; absence = skill path not used

    system_prompt = (
        "You are an expert technical resume writer and ATS optimization specialist with the "
        "`docx` skill available for producing Word documents.\n"
        "CRITICAL RULES:\n"
        "1. PRESERVE every substantive point from the candidate's existing resume — never drop "
        "or weaken their real experience. Include each original point as an item with "
        "is_new=false (light wording cleanup is fine, but keep the original meaning).\n"
        "2. ADD new, TRUTHFUL, JD-tailored bullet points and keywords that strengthen the match. "
        "Mark every point you add with is_new=true. Do not fabricate employers or fake metrics.\n"
        "3. Organize into a `summary` (Professional Summary paragraph) and `sections` (e.g. Core "
        "Skills, Professional Experience, Certifications). Each section has a title and items; "
        "each item has `text` and `is_new`.\n"
        f"4. Using the docx skill, build a clean, professional, ATS-friendly Word document of the "
        f"optimized resume and save it to the absolute path: {out_path} . Write ONLY to that "
        "path — do not create or modify any other files. The downloaded resume must be clean "
        "(do NOT include diff markers/highlighting in the document itself).\n"
        "Your FINAL message MUST be ONLY the JSON object matching this schema (no prose, no code "
        f"fences):\n{json.dumps(ResumeContent.model_json_schema())}"
    )
    prompt = (
        f"JOB DESCRIPTION / REQUIREMENT:\n{job_description}\n\n"
        f"CANDIDATE'S EXISTING RESUME:\n{resume_text or '(no resume provided)'}\n\n"
        "Preserve all original points (is_new=false) and add tailored new points (is_new=true). "
        f"Build the polished .docx at {out_path} using the docx skill, then return the structured "
        "JSON as your final message."
    )
    options = ClaudeAgentOptions(
        model=None,
        cwd=REPO_ROOT,                       # so the SDK discovers .claude/skills/
        setting_sources=["project"],         # load project skills (the bundled docx skill)
        skills=["docx"],                     # enable the docx skill (adds the Skill tool)
        max_turns=40,
        allowed_tools=["Read", "Write", "Edit", "Bash", "Glob", "Grep", "Skill"],
        permission_mode="bypassPermissions",
        system_prompt=system_prompt,
        output_format=ResumeContent.model_json_schema(),
    )

    result_text = ""
    structured = None
    try:
        async with ClaudeSDKClient(options) as client:
            await client.query(prompt)
            async for msg in client.receive_response():
                if isinstance(msg, AssistantMessage):
                    for block in msg.content:
                        if isinstance(block, TextBlock):
                            result_text += block.text
                elif isinstance(msg, ResultMessage):
                    if msg.is_error:
                        raise RuntimeError(msg.errors or msg.result or "Claude returned an error")
                    if msg.structured_output:
                        structured = msg.structured_output
                    elif msg.result and not result_text:
                        result_text = msg.result
                    break

        content = (
            sanitize_content(ResumeContent(**structured).model_dump())
            if structured
            else _parse_content(result_text)
        )

        # Use the skill-produced document when present; otherwise the caller falls back.
        docx_bytes = None
        if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
            with open(out_path, "rb") as f:
                docx_bytes = f.read()
        return content, docx_bytes
    finally:
        if os.path.exists(out_path):
            try:
                os.remove(out_path)
            except OSError:
                pass


async def run_resume_optimization(user_id: int, job_description: str, resume_text: str):
    try:
        _set_job(user_id, status="running", stage="Parsing resume", progress=20)
        _set_job(user_id, stage="Optimizing with Claude + docx skill", progress=45)
        content, docx_bytes = await _optimize_with_claude(job_description, resume_text)
        _set_job(user_id, stage="Building Word document", progress=80)
        markdown = content_to_markdown(content)
        # Prefer the skill-built .docx; fall back to a deterministic python-docx build.
        if not docx_bytes:
            docx_bytes = build_docx_from_content(content)
        _set_job(
            user_id,
            status="done",
            stage="Complete",
            progress=100,
            result_json=json.dumps(content),
            result_markdown=markdown,
            result_docx=docx_bytes,
            error="",
        )
    except Exception as e:  # noqa: BLE001
        _set_job(user_id, status="error", stage="Failed", error=str(e))


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------
router = APIRouter(prefix="/api/resume", tags=["resume"])


@router.post("/optimize")
async def optimize(
    background_tasks: BackgroundTasks,
    job_description: str = Form(...),
    original_text: str | None = Form(None),
    resume: UploadFile | None = File(None),
    user: dict = Depends(get_current_user),
):
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description is required.")

    # Precedence: edited text from the UI > a freshly uploaded .docx > the stored original.
    resume_text = (original_text or "").strip()
    if not resume_text and resume is not None:
        data = await resume.read()
        if not resume.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Please upload a .docx Word file.")
        try:
            resume_text = extract_text_from_docx(data)
        except Exception:
            raise HTTPException(status_code=400, detail="Could not read the Word document.")

    if not resume_text:
        existing = _get_job(user["id"])
        if existing is not None:
            resume_text = existing["original_text"] or ""

    _set_job(
        user["id"],
        status="running",
        stage="Queued",
        progress=5,
        job_description=job_description,
        original_text=resume_text,
        result_markdown="",
        result_json="",
        result_docx=None,
        error="",
    )
    background_tasks.add_task(
        run_resume_optimization, user["id"], job_description, resume_text
    )
    return {"job_id": user["id"], "status": "running"}


@router.get("/status")
async def status(user: dict = Depends(get_current_user)):
    row = _get_job(user["id"])
    if row is None:
        return {"status": "idle", "stage": "", "progress": 0, "error": "", "has_result": False}
    return {
        "status": row["status"],
        "stage": row["stage"],
        "progress": row["progress"],
        "error": row["error"] or "",
        "has_result": bool(row["result_json"] or row["result_markdown"]),
        "job_description": row["job_description"] or "",
    }


@router.get("/result")
async def result(user: dict = Depends(get_current_user)):
    row = _get_job(user["id"])
    if row is None or not (row["result_json"] or row["result_markdown"]):
        raise HTTPException(status_code=404, detail="No optimized resume available yet.")
    try:
        content = json.loads(row["result_json"]) if row["result_json"] else None
    except Exception:
        content = None
    return {
        "content": content,
        "markdown": row["result_markdown"],
        "original_text": row["original_text"] or "",
        "stage": row["stage"],
        "status": row["status"],
    }


@router.put("/content")
async def save_content(
    content: ResumeContent = Body(..., embed=True),
    user: dict = Depends(get_current_user),
):
    """Save user-edited optimized content and rebuild the downloadable .docx."""
    row = _get_job(user["id"])
    if row is None:
        raise HTTPException(status_code=404, detail="No resume job to update.")
    data = sanitize_content(content.model_dump())
    markdown = content_to_markdown(data)
    docx_bytes = build_docx_from_content(data)
    _set_job(
        user["id"],
        status="done",
        result_json=json.dumps(data),
        result_markdown=markdown,
        result_docx=docx_bytes,
    )
    return {"success": True}


@router.get("/download")
async def download(user: dict = Depends(get_current_user)):
    row = _get_job(user["id"])
    if row is None or row["result_docx"] is None:
        raise HTTPException(status_code=404, detail="No generated document available yet.")
    return StreamingResponse(
        io.BytesIO(row["result_docx"]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="optimized_resume.docx"'},
    )
